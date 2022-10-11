from tkinter import *
import tkinter
import threading
import tkinter as frame
from tkinter import messagebox
from tkinter import ttk
from tkinter.ttk import Notebook, Style
from tkinter.scrolledtext import ScrolledText
from PIL import Image
import pytesseract
from googletrans import Translator
import os
import pandas as pd
from docx import Document
import re
from datetime import date
from docx.shared import Pt
from docx.shared import Inches
import pyautogui
import glob
import shutil
from selenium import webdriver
import selenium.webdriver.chrome.options
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait as Wait
from selenium.webdriver.support import expected_conditions as ec
from selenium.common.exceptions import UnexpectedAlertPresentException
from selenium.common.exceptions import InvalidArgumentException, NoSuchWindowException, WebDriverException
from selenium.common.exceptions import NoSuchElementException, ElementClickInterceptedException
from selenium.webdriver.chrome.service import Service
import urllib
import threading
from restrict.urls import Urls

Translator = Translator()


def bo_sf_china():
    log_box.config(state='normal')
    log_box.delete("1.0", END)
    log_box.insert(END, "Browser initialization....\n")
    log_box.config(state='disabled')
    first_name = str(first_name2.get())
    last_name = str(last_name2.get())
    client_id = str(client_id2.get())
    email = str(email2.get())
    salesforce_driver = webdriver.Chrome(service=ser, options=driver_options)
    salesforce_driver.set_window_size(1920, 1080)
    driver = webdriver.Chrome(service=ser, options=driver_options)
    driver.set_window_size(1920, 1080)

    def login_to_BO():
        try:
            form_url = Urls.profile_url + str(client_id) + "/edit"
            driver.get(form_url)
            username = driver.find_element(By.NAME, "admin_username")
            username.send_keys("")
            password = driver.find_element(By.NAME, "admin_password")
            password.send_keys("")
            driver.find_element(By.CLASS_NAME, "col-xs-4").click()
            log_box.config(state='normal')
            log_box.insert(END, "Successfully logged in....\n")
            log_box.config(state='disabled')
        except NoSuchElementException:
            login_to_BO()
            log_box.config(state='normal')
            log_box.insert(END, "Login failed...retrying...\n")
            log_box.config(state='disabled')
        log_box.config(state='normal')
        log_box.insert(END, "Logging to Backoffice....\n")
        log_box.config(state='disabled')

    def open_salesforce():
        salesforce_url = "https://login.salesforce.com/"
        salesforce_driver.get(salesforce_url)
        username = salesforce_driver.find_element(By.NAME, "username")
        username.send_keys("")
        password = salesforce_driver.find_element(By.NAME, "pw")
        password.send_keys("")
        salesforce_driver.find_element(By.NAME, "Login").click()
        salesforce_driver.find_element(By.XPATH, '//*[@id="tryLexDialogX"]').click()
        salesforce_driver.find_element(By.ID, "phSearchInput").send_keys(email)
        salesforce_driver.find_element(By.ID, "phSearchButton").click()
        time.sleep(4)
        salesforce_driver.find_element(By.XPATH, '//*[@id="Contact_body"]/table/tbody/tr[2]/td[1]/a').click()
        time.sleep(2)
        salesforce_driver.find_element(By.ID, "00N0I00000JsRdF").clear()
        try:
            salesforce_driver.find_element(By.ID, "00N0I00000JsRdF").send_keys(str(full_name))
            salesforce_driver.find_element(By.NAME, 'save').click()
            salesforce_driver.find_element(By.XPATH, '//*[@id="Account_body"]/table/tbody/tr[2]/th/a').click()
        except NameError:
            log_box.config(state='normal')
            log_box.insert(END, "Waiting for updating the Local Name on salesforce...\n")
            log_box.config(state='disabled')

    def change_name():
        log_box.config(state='normal')
        log_box.insert(END, "Updating client name...\n")
        log_box.config(state='disabled')
        firstname = driver.find_element(By.ID, "s2276c242cd_PortalUser_first_name")
        firstname.clear()
        firstname.send_keys(str(first_name))
        lastname = driver.find_element(By.ID, "s2276c242cd_PortalUser_last_name")
        lastname.clear()
        lastname.send_keys(str(last_name))

    def personal_information():
        try:
            log_box.config(state='normal')
            log_box.insert(END, "Updating personal information...\n")
            log_box.config(state='disabled')
            click_personal_information = driver.find_element(By.LINK_TEXT, "Personal Information (2)")
            driver.execute_script("$(arguments[0]).click();", click_personal_information)
            driver.execute_script("window.scrollTo(0, 1100);")
            Wait(driver, 10).until(ec.visibility_of_element_located((By.XPATH, '//*[@id="select2-chosen-22"]')))
            driver.find_element(By.ID, "s2276c242cd_employmentInformation__businessSector").clear()
            driver.find_element(By.ID, "s2276c242cd_employmentInformation__businessSector").send_keys("Trading")
            driver.find_element(By.ID, "s2276c242cd_employmentInformation__occupation").clear()
            driver.find_element(By.ID, "s2276c242cd_employmentInformation__occupation").send_keys("Manager")
            Wait(driver, 10).until(ec.element_to_be_clickable((By.XPATH, '//span[contains(@id,"chosen-22")]'))).click()
            pyautogui.moveTo(x=365, y=773)
            pyautogui.click()
        except (UnexpectedAlertPresentException, ElementClickInterceptedException):
            log_box.config(state='normal')
            log_box.insert(END, "Updating failed...don't move your mouse...retrying\n")
            log_box.config(state='disabled')
            personal_information()

    def trading_info():
        try:
            log_box.config(state='normal')
            log_box.insert(END, "Updating trading info...\n")
            log_box.config(state='disabled')
            click_personal_information = driver.find_element(By.LINK_TEXT, "Trading Info.")
            driver.execute_script("$(arguments[0]).click();", click_personal_information)
            driver.execute_script("window.scrollTo(0, 1100);")
            Wait(driver, 10).until(ec.element_to_be_clickable((By.XPATH, '//span[contains(@id,"chosen-37")]'))).click()
            pyautogui.moveTo(x=1200, y=890)
            time.sleep(1)
            pyautogui.click()
            driver.execute_script("window.scrollTo(0,  300);")
            Wait(driver, 10).until(ec.element_to_be_clickable((By.NAME, 'btn_update_and_edit'))).click()
        except (UnexpectedAlertPresentException, ElementClickInterceptedException):
            log_box.config(state='normal')
            log_box.insert(END, "Updating failed...don't move your mouse...retrying\n")
            log_box.config(state='disabled')
            trading_info()

    def generate_pdf():
        # 跳转页码
        log_box.config(state='normal')
        log_box.insert(END, "Generating application form...\n")
        log_box.config(state='disabled')
        client_id1 = str(client_id2.get())
        bo_show_url = Urls.profile_url + str(client_id1) + "/show"
        driver.get(bo_show_url)
        click_documents = driver.find_element(By.LINK_TEXT, "Documents")
        driver.execute_script("$(arguments[0]).click();", click_documents)
        time.sleep(1)
        # 生成PDF
        generate_button = driver.find_element(By.LINK_TEXT, "Generate pdf")
        generate_button.click()
        pdf_file = driver.find_element(By.LINK_TEXT, "PDF FILE")
        pdf_file.click()

    def copy_paste_application_form():
        log_box.config(state='normal')
        log_box.insert(END, "Preparing application form...\n")
        log_box.config(state='disabled')
        list_of_files = glob.glob('Downloads\*.pdf')
        latest_file = max(list_of_files, key=os.path.getctime)
        original = latest_file
        try:
            target = Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_APP.pdf"
            shutil.copyfile(original, target)
        except FileNotFoundError:
            os.makedirs(Urls.os_path + str(email))
            target = Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_APP.pdf"
            shutil.copyfile(original, target)

    def download_file():
        log_box.config(state='normal')
        log_box.insert(END, "Downloading ID...\n")
        log_box.config(state='disabled')
        id_img = driver.find_element(By.XPATH, '//img[contains(@class,"img-fluid width2")]')
        src = id_img.get_attribute('src')
        urllib.request.urlretrieve(src, Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_idfront.jpg")

    def upload_file():
        driver.find_element(By.ID, "dZUploadLPOA1-upload")
        log_box.config(state='normal')
        log_box.insert(END, "Uploading ID/COT/APP...\n")
        log_box.config(state='disabled')
        try:
            driver.find_element(By.XPATH, '/html/body/input[3]').send_keys(Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_idfront.jpg")
            time.sleep(1)
            driver.find_element(By.XPATH, '/html/body/input[3]').send_keys(Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_COT.pdf")
            time.sleep(1)
            driver.find_element(By.XPATH, '/html/body/input[3]').send_keys(Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_APP.pdf")
        except InvalidArgumentException:
            log_box.config(state='normal')
            log_box.insert(END, "Unable to upload the file...check availability...\n")
            log_box.config(state='disabled')
            tkinter.messagebox.showerror(title='Warning', message='Unable to upload the file, check availability!')

    def approve_id_and_por():
        approve_id_url = Urls.profile_url + str(client_id) + "/show"
        driver.get(approve_id_url)
        log_box.config(state='normal')
        log_box.insert(END, "Checking approval possibility...\n")
        log_box.config(state='disabled')
        try:
            driver.find_element(By.LINK_TEXT, Urls.manager_url + "approve_id/" + str(client_id)).click()
            driver.get(approve_id_url)
            log_box.config(state='normal')
            log_box.insert(END, "ID is approved...\n")
            log_box.config(state='disabled')
        except NoSuchElementException:
            log_box.config(state='normal')
            log_box.insert(END, "ID maybe already approved...check ID status...\n")
            log_box.config(state='disabled')
        try:
            driver.find_element(By.LINK_TEXT, Urls.manager_url + "approve_por/" + str(client_id)).click()
            log_box.config(state='normal')
            log_box.insert(END, "POR is approved...\n")
            log_box.config(state='disabled')

        except NoSuchElementException:
            log_box.config(state='normal')
            log_box.insert(END, "POR maybe already approved...check POR status...\n")
            log_box.config(state='disabled')

    login_to_BO()
    change_name()
    personal_information()
    trading_info()
    generate_pdf()
    time.sleep(3)
    copy_paste_application_form()
    download_file()
    upload_file()
    approve_id_and_por()
    open_salesforce()


def bo_sf_other():
    log_box.config(state='normal')
    log_box.delete("1.0", END)
    log_box.insert(END, "Browser initialization....\n")
    log_box.config(state='disabled')
    first_name = str(first_name2.get())
    last_name = str(last_name2.get())
    client_id = str(client_id2.get())
    email = str(email2.get())
    try:
        os.makedirs(Urls.os_path + str(email))
    except OSError:
        tkinter.messagebox.showerror(title='Warning', message='The folder is already exist!')
    salesforce_driver = webdriver.Chrome(service=ser, options=driver_options)
    salesforce_driver.set_window_size(1920, 1080)
    driver = webdriver.Chrome(service=ser, options=driver_options)
    driver.set_window_size(1920, 1080)

    def login_to_BO():
        log_box.config(state='normal')
        log_box.insert(END, "Logging to Backoffice....\n")
        log_box.config(state='disabled')
        try:
            form_url = Urls.profile_url + str(client_id) + "/edit"
            driver.get(form_url)
            username = driver.find_element(By.NAME, "admin_username")
            username.send_keys("")
            password = driver.find_element(By.NAME, "admin_password")
            password.send_keys("")
            driver.find_element(By.CLASS_NAME, "col-xs-4").click()
        except NoSuchElementException:
            login_to_BO()
            log_box.config(state='normal')
            log_box.insert(END, "Login failed...retrying...\n")
            log_box.config(state='disabled')

    def open_salesforce():
        salesforce_url = "https://login.salesforce.com/"
        salesforce_driver.get(salesforce_url)
        username = salesforce_driver.find_element(By.NAME, "username")
        username.send_keys("")
        password = salesforce_driver.find_element(By.NAME, "pw")
        password.send_keys("")
        salesforce_driver.find_element(By.NAME, "Login").click()
        salesforce_driver.find_element(By.XPATH, '//*[@id="tryLexDialogX"]').click()
        salesforce_driver.find_element(By.ID, "phSearchInput").send_keys(email)
        salesforce_driver.find_element(By.ID, "phSearchButton").click()
        time.sleep(4)
        salesforce_driver.find_element(By.XPATH, '//*[@id="Account_body"]/table/tbody/tr[2]/th/a').click()

    def change_name():
        log_box.config(state='normal')
        log_box.insert(END, "Updating client name...\n")
        log_box.config(state='disabled')
        firstname = driver.find_element(By.ID, "s2276c242cd_PortalUser_first_name")
        firstname.clear()
        firstname.send_keys(str(first_name))
        lastname = driver.find_element(By.ID, "s2276c242cd_PortalUser_last_name")
        lastname.clear()
        lastname.send_keys(str(last_name))

    def personal_information():
        try:
            log_box.config(state='normal')
            log_box.insert(END, "Updating personal information...\n")
            log_box.config(state='disabled')
            click_personal_information = driver.find_element(By.LINK_TEXT, "Personal Information (2)")
            driver.execute_script("$(arguments[0]).click();", click_personal_information)
            driver.execute_script("window.scrollTo(0, 1100);")
            Wait(driver, 10).until(ec.visibility_of_element_located((By.XPATH, '//*[@id="select2-chosen-22"]')))
            driver.find_element(By.ID, "s2276c242cd_employmentInformation__businessSector").clear()
            driver.find_element(By.ID, "s2276c242cd_employmentInformation__businessSector").send_keys("Trading")
            driver.find_element(By.ID, "s2276c242cd_employmentInformation__occupation").clear()
            driver.find_element(By.ID, "s2276c242cd_employmentInformation__occupation").send_keys("Manager")
            Wait(driver, 10).until(ec.element_to_be_clickable((By.XPATH, '//span[contains(@id,"chosen-22")]'))).click()
            pyautogui.moveTo(x=365, y=773)
            pyautogui.click()
        except (UnexpectedAlertPresentException, ElementClickInterceptedException):
            log_box.config(state='normal')
            log_box.insert(END, "Updating failed...don't move your mouse...retrying\n")
            log_box.config(state='disabled')
            personal_information()

    def trading_info():
        try:
            log_box.config(state='normal')
            log_box.insert(END, "Updating trading info...\n")
            log_box.config(state='disabled')
            click_personal_information = driver.find_element(By.LINK_TEXT, "Trading Info.")
            driver.execute_script("$(arguments[0]).click();", click_personal_information)
            driver.execute_script("window.scrollTo(0, 1100);")
            Wait(driver, 10).until(ec.element_to_be_clickable((By.XPATH, '//span[contains(@id,"chosen-37")]'))).click()
            pyautogui.moveTo(x=1200, y=890)
            time.sleep(1)
            pyautogui.click()
            driver.execute_script("window.scrollTo(0,  300);")
            Wait(driver, 10).until(ec.element_to_be_clickable((By.NAME, 'btn_update_and_edit'))).click()
        except (UnexpectedAlertPresentException, ElementClickInterceptedException):
            log_box.config(state='normal')
            log_box.insert(END, "Updating failed...don't move your mouse...retrying\n")
            log_box.config(state='disabled')
            trading_info()

    def generate_pdf():
        # 跳转页码
        log_box.config(state='normal')
        log_box.insert(END, "Generating application form...\n")
        log_box.config(state='disabled')
        client_id1 = str(client_id2.get())
        bo_show_url = Urls.profile_url + str(client_id1) + "/show"
        driver.get(bo_show_url)
        click_documents = driver.find_element(By.LINK_TEXT, "Documents")
        driver.execute_script("$(arguments[0]).click();", click_documents)
        time.sleep(1)
        # 生成PDF
        generate_button = driver.find_element(By.LINK_TEXT, "Generate pdf")
        generate_button.click()
        pdf_file = driver.find_element(By.LINK_TEXT, "PDF FILE")
        pdf_file.click()

    def copy_paste_application_form():
        log_box.config(state='normal')
        log_box.insert(END, "Preparing application form...\n")
        log_box.config(state='disabled')
        list_of_files = glob.glob('Downloads\*.pdf')
        latest_file = max(list_of_files, key=os.path.getctime)
        original = latest_file
        target = Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_APP.pdf"
        shutil.copyfile(original, target)

    def download_and_upload_the_file():
        log_box.config(state='normal')
        log_box.insert(END, "Uploading application form...\n")
        log_box.config(state='disabled')
        driver.find_element(By.ID, "dZUploadLPOA1-upload")
        try:
            time.sleep(1)
            driver.find_element(By.XPATH, '/html/body/input[4]').send_keys(Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_APP.pdf")
        except InvalidArgumentException:
            log_box.config(state='normal')
            log_box.insert(END, "Unable to upload the file...check availability...\n")
            log_box.config(state='disabled')
            tkinter.messagebox.showerror(title='Warning', message='Unable to upload the file, check availability!')

    def approve_id_and_por():
        approve_id_url = "view-source:" + Urls.profile_url + str(client_id) + "/show"
        driver.get(approve_id_url)
        log_box.config(state='normal')
        log_box.insert(END, "Checking approval possibility...\n")
        log_box.config(state='disabled')
        try:
            driver.find_element(By.LINK_TEXT, Urls.manager_url + "approve_id/" + str(client_id)).click()
            driver.get(approve_id_url)
            log_box.config(state='normal')
            log_box.insert(END, "ID is approved...\n")
            log_box.config(state='disabled')
        except NoSuchElementException:
            log_box.config(state='normal')
            log_box.insert(END, "ID maybe already approved...check ID status...\n")
            log_box.config(state='disabled')
        try:
            driver.find_element(By.LINK_TEXT, Urls.manager_url + "approve_por/" + str(client_id)).click()
            log_box.config(state='normal')
            log_box.insert(END, "POR is approved...\n")
            log_box.config(state='disabled')
        except NoSuchElementException:
            log_box.config(state='normal')
            log_box.insert(END, "POR maybe already approved...check POR status...\n")
            log_box.config(state='disabled')

    login_to_BO()
    change_name()
    personal_information()
    trading_info()
    generate_pdf()
    time.sleep(3)
    copy_paste_application_form()
    download_and_upload_the_file()
    approve_id_and_por()
    open_salesforce()


def id_translation():
    document = Document()
    pytesseract.pytesseract.tesseract_cmd = r"C:\Tesseract-OCR\tesseract.exe"
    email = str(email2.get())
    first_name = str(first_name2.get())
    last_name = str(last_name2.get())
    try:
        os.makedirs(Urls.os_path + str(email))
    except OSError:
        tkinter.messagebox.showerror(title='Warning', message='The folder is already exist!')

    document.save(Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_COT.docx")
    # 翻译反面
    path = "ID_4.jpg"
    id_back = str(pytesseract.image_to_string(Image.open(path), lang='chi_sim')).replace(" ", "").replace(".", "")
    id_back = str(re.findall('\d', id_back)).replace(", ", "").replace("'", "")
    id_back = id_back.replace(" ", "").replace("[", "")
    issue_date = id_back[:4] + "-" + id_back[4:6] + "-" + id_back[6:8]
    expiry_date = id_back[8:12] + "-" + id_back[4:6] + "-" + id_back[6:8]

    # 翻译正面
    path = "ID_3.jpg"
    id_front_zh = pytesseract.image_to_string(Image.open(path), lang='chi_sim').replace("\n", "").replace('"', "")
    str(Translator.translate(id_front_zh, src="zh-CN", dest="en"))

    # 处理正面
    id_front_zh = pytesseract.image_to_string(Image.open(path), lang='chi_sim')
    id_front_zh = id_front_zh.replace(" ", "")
    # full_name = id_front_zh.splitlines()[0]
    # full_name_en = str(Translator.translate(full_name))
    # 居住地址
    id_front_zh = pytesseract.image_to_string(Image.open(path), lang='chi_sim').replace("\n", "").replace(" ", "").strip()
    Address_zh = str(re.findall(r'日(.*?)$', id_front_zh)).replace("[", "").replace("]", "")
    if len(Address_zh) == 0:
        Address_zh = str(re.findall(r'目(.*?)$', id_front_zh)).replace("[", "").replace("]", "")
    if len(Address_zh) == 0:
        Address_zh = str(re.findall(r'晶(.*?)$', id_front_zh)).replace("[", "").replace("]", "")
    Address_en = str(Translator.translate(Address_zh, src="zh-CN", dest="en"))
    str(Translator.translate(id_front_zh, src="zh-CN", dest="en"))

    # 生日
    DOB = "1" + str(re.findall(r'1(.*?)日', id_front_zh)).replace("年", "-").replace("月", "-")
    DOB = str(DOB.replace("第", "-"))
    DOB = str(DOB.replace(" ", "").replace("'", ""))
    DOB = str(DOB.replace("[", "").replace("]", ""))
    if DOB[:2] != '19':
        DOB = "2" + str(re.findall(r'2(.*?)日', id_front_zh)).replace("年", "-").replace("月", "-")
        DOB = str(DOB.replace("第", "-"))
        DOB = str(DOB.replace(" ", "").replace("'", ""))
        DOB = str(DOB.replace("[", "").replace("]", ""))
    if DOB[:2] == "2":
        DOB = "1" + str(re.findall(r'1(.*?)目', id_front_zh)).replace("年", "-").replace("月", "-")
        DOB = str(DOB.replace("第", "-"))
    if DOB[:2] == '2':
        DOB = "1" + str(re.findall(r'1(.*?)卓', id_front_zh)).replace("年", "-").replace("月", "-")
    if DOB[:2] == '2':
        DOB = "1" + str(re.findall(r'1(.*?)晶', id_front_zh)).replace("年", "-").replace("月", "-")
        DOB = str(DOB.replace(" ", "").replace("'", ""))
        DOB = str(DOB.replace("[", "").replace("]", ""))
    # Date
    Today = str(date.today())
    # 符号去除
    disallowed_char = ", '[]:"

    for char in disallowed_char:
        issue_date = issue_date.replace(disallowed_char, "").replace('"', '')
        expiry_date = expiry_date.replace(disallowed_char, "").replace('"', '')
        Address_zh = Address_zh.replace(disallowed_char, "").replace('"', '')
        Address_en = Address_en.replace(disallowed_char, "").replace('"', '')
        # full_name = full_name.replace(disallowed_char, "").replace('"', '')
        # full_name_en = full_name_en.replace(disallowed_char, "").replace('"', '')
        DOB = DOB.replace(char, "").replace('"', '')

    issue_date = issue_date.replace("Translated(src=zh-CN, dest=en, text=", "").replace(", pronunciation=None, extra_data={'translat...)", '')
    expiry_date = expiry_date.replace("Translated(src=zh-CN, dest=en, text=", "").replace(", pronunciation=None, extra_data={'translat...)", '')
    Address_zh.replace("Translated(src=zh-CN, dest=en, text=", "").replace(", pronunciation=None, extra_data={'translat...)", '')
    Address_en = Address_en.replace("Translated(src=zh-CN, dest=en, text=", "").replace(", pronunciation=None, extra_data={'translat...)", '')
    Address_en = Address_en.strip()
    # full_name = full_name.replace("Translated(src=zh-CN, dest=en, text=", "").replace(", pronunciation=None, extra_data={'translat...)", '')
    # full_name_en = full_name_en.replace("Translated(src=zh-CN, dest=en, text=", "").replace(", pronunciation=None, extra_data={'translat...)", '')
    DOB = DOB.replace("Translated(src=zh-CN, dest=en, text=", "").replace(", pronunciation=None, extra_data={'translat...)", '')

    # 格式
    heading1 = document.add_heading("Confirmation of Translation")
    heading1.alignment = 1
    document.add_paragraph("I, Charmy as Account Application Officer confirm that I speak and read comprehensively and fluently in Chinese and English language.")
    document.add_paragraph("I refer to the attachments to these documents, namely ID and Proof of Residence. ")
    document.add_paragraph("I confirm to the best of my knowledge and belief that the following translation below (as to the highlighted parts in the attachment) are accurate.")
    document.add_heading("Identification")
    doc = document.add_paragraph("Identification Type: ID").add_run()
    doc.font.name = 'Times New Roman'
    style = document.styles['Normal']
    style.paragraph_format.space_after = Pt(4)
    document.add_paragraph("Full name: " + str(first_name) + ' ' + str(last_name))
    document.add_paragraph("DOB: " + DOB)
    document.add_paragraph("Expiry date:" + str(expiry_date))
    document.add_paragraph("Nationality: Chinese")
    document.add_heading("Proof of residence")
    document.add_paragraph("Document Type: ID")
    document.add_paragraph("Full name: " + str(first_name) + ' ' + str(last_name))
    document.add_paragraph("Address: " + str(Address_en).replace("'", ""))
    document.add_paragraph("Issue date: " + str(issue_date))
    document.add_paragraph("Signed:")
    sign = document.add_paragraph()
    sign1 = sign.add_run("                     Charmy")
    sign1.font.name = 'Bradley Hand ITC'
    document.add_paragraph("Name: Charmy")
    document.add_paragraph("Date: " + Today)
    document.save(Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_COT.docx")
    os.startfile(Urls.os_path + str(email) + "\\" + str(first_name) + str(last_name) + "_COT.docx")


def radiobutton_option():
    choice = variable.get()
    if choice == "China":
        try:
            bo_sf_china()
            log_box.config(state='normal')
            log_box.insert(END, "Application process succeed ...\n")
            log_box.config(state='disabled')
        except (NoSuchWindowException, WebDriverException):
            log_box.config(state='normal')
            log_box.insert(END, "Browser has been terminated...application process failed...\n")
            log_box.config(state='disabled')
    elif choice == "Other":
        try:
            bo_sf_other()
            log_box.config(state='normal')
            log_box.insert(END, "Application process succeed ...\n")
            log_box.config(state='disabled')
        except (NoSuchWindowException, WebDriverException):
            log_box.config(state='normal')
            log_box.insert(END, "Browser has been terminated...application process failed...\n")
            log_box.config(state='disabled')
    else:
        tkinter.messagebox.showerror(title='Warning', message='Please select the country!')


def show_client_id():
    read_file = pd.read_excel('monitor.xlsx')
    read_file.to_csv('monitor.csv', index=None)
    client_id2.delete(0, END)
    input_file = 'monitor.csv'
    show_id_df = pd.read_csv(input_file)
    email = str(email2.get())
    for i in range(len(show_id_df)):
        if show_id_df.iloc[i, 1] == email:
            client_id2.insert(END, show_id_df.iloc[i, 0])
            break


def vist_bo():
    log_box.config(state='normal')
    log_box.delete("1.0", END)
    log_box.insert(END, "Browser initialization....\n")
    log_box.config(state='disabled')
    first_name = str(first_name2.get())
    last_name = str(last_name2.get())
    client_id = str(client_id2.get())
    email = str(email2.get())
    driver = webdriver.Chrome(service=ser, options=driver_options)
    driver.set_window_size(1920, 1080)

    def login_to_BO():
        try:
            form_url = Urls.profile_url + str(client_id) + "/show"
            driver.get(form_url)
            username = driver.find_element(By.NAME, "admin_username")
            username.send_keys("")
            password = driver.find_element(By.NAME, "admin_password")
            password.send_keys("")
            driver.find_element(By.CLASS_NAME, "col-xs-4").click()
            log_box.config(state='normal')
            log_box.insert(END, "Successfully logged in....\n")
            log_box.config(state='disabled')
        except NoSuchElementException:
            login_to_BO()
            log_box.config(state='normal')
            log_box.insert(END, "Login failed...retrying...\n")
            log_box.config(state='disabled')
    login_to_BO()


def show_monitor():
    os.startfile("dist/Schedule.exe")


global full_name
global full_name_en

root = frame.Tk()
variable = StringVar(root)
variable.set("Type")
root.resizable(False, False)
root['background'] = '#000000'
ser = Service('chromedriver.exe')
driver_options = webdriver.ChromeOptions()
driver_options.add_experimental_option("detach", True)
driver_options.add_experimental_option('excludeSwitches', ['enable-logging'])
root.iconbitmap(r'1.ico')
root.title("NINJA 0.1.3")
root.geometry("430x510")
window_width = 430
window_height = 510
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
# find the center point
center_x = int(screen_width / 2 - window_width / 2)
center_y = int(screen_height / 2 - window_height / 2)
# set the position of the window to the center of the screen
root.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')


style = ttk.Style()

style.theme_use('vista')

style.configure('TLabel',
                background="black",
                foreground="#e6b04e",
                fieldbackground="black",
                font=('Times New Roman', 10),
                )

ttk.Label(root,
          text="Client ID",
          ).place(x=25, y=60)

ttk.Label(root,
          text="First Name"
          ).place(x=25, y=90)

ttk.Label(root,
          text="Last Name",
          style="BW.TLabel"
          ).place(x=25, y=120)

ttk.Label(root,
          text="Email",
          style='TLabel',
          ).place(x=25, y=30)

# Entry
email2 = Entry(root,
               bg='#000000',
               fg='#e6b04e',
               font=('Times New Roman', 10),
               )

client_id2 = Entry(root,
                   bg='#000000',
                   fg='#e6b04e',
                   font=('Times New Roman', 10),
                   )

first_name2 = Entry(root,
                    bg='#000000',
                    fg='#e6b04e',
                    font=('Times New Roman', 10),
                    )

last_name2 = Entry(root,
                   bg='#000000',
                   fg='#e6b04e',
                   font=('Times New Roman', 10),
                   )

log_box = ScrolledText(root,
                       height=17,
                       width=71,
                       background='#000000',
                       foreground='#e6b04e',
                       font=('Times New Roman', 12),
                       state='disabled',
                       )

# Button
frame.Button(root,
             bg='#000000',
             fg='#e6b04e',
             text="ID TRANS",
             font=('Times New Roman', 9),
             command=lambda: threading.Thread(target=id_translation).start(),
             ).place(x=270, y=15, height=30, width=80)


frame.Button(root,
             bg='#000000',
             fg='#e6b04e',
             font=('Times New Roman', 9),
             text="SHOW ID",
             command=show_client_id,
             ).place(x=270, y=45, height=30, width=80)

frame.Button(root,
             bg='#000000',
             fg='#e6b04e',
             text="BO & SF",
             font=('Times New Roman', 9),
             command=lambda: threading.Thread(target=radiobutton_option).start(),
             ).place(x=270, y=75, height=30, width=80)

frame.Button(root,
             bg='#000000',
             fg='#e6b04e',
             text="VIST BO",
             font=('Times New Roman', 9),
             command=lambda: threading.Thread(target=vist_bo).start(),
             ).place(x=270, y=105, height=30, width=80)

frame.Button(root,
             bg='#000000',
             fg='#e6b04e',
             text="MONITOR",
             font=('Times New Roman', 9),
             command=show_monitor,
             ).place(x=270, y=135, height=30, width=80)

Radiobutton(root,
            bg='#000000',
            fg='#e6b04e',
            relief='flat',
            text='China',
            value='China',
            font=('Times New Roman', 10),
            variable=variable,
            ).place(x=360, y=70)

Radiobutton(root,
            bg='#000000',
            fg='#e6b04e',
            relief='flat',
            text='Other',
            font=('Times New Roman', 10),
            value='Other',
            variable=variable,
            ).place(x=360, y=90)


# Label place
first_name2.place(x=90,
                  y=85,
                  height=25,
                  width=130,
                  )

last_name2.place(x=90,
                 y=115,
                 height=25,
                 width=130,
                 )

email2.place(x=90,
             y=25,
             height=25,
             width=130,
             )

client_id2.place(x=90,
                 y=55,
                 height=25,
                 width=130,
                 )

log_box.place(x=0,
              y=180,
              )

root.mainloop()





